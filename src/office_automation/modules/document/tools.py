"""文档模块 — 核心工具函数

使用 python-docx 操作 Word .docx 文档
支持模板占位符替换和 CSV 数据源批量生成
"""
import csv
import json
from pathlib import Path
from office_automation.shared.utils import output_path, ensure_dir, timestamp


def _load_data_source(source_path: str) -> list[dict]:
    """加载数据源（CSV 或 JSON），返回 dict 列表"""
    p = Path(source_path)
    if not p.exists():
        raise FileNotFoundError(f"数据源文件不存在: {source_path}")

    if p.suffix.lower() == ".csv":
        rows = []
        with open(p, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(dict(row))
        return rows
    elif p.suffix.lower() == ".json":
        data = json.loads(p.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else [data]
    else:
        raise ValueError(f"不支持的数据源格式: {p.suffix}，请使用 .csv 或 .json")


def _replace_in_paragraphs(doc, replace_map: dict):
    """替换文档中所有段落里的占位符"""
    for para in doc.paragraphs:
        for key, val in replace_map.items():
            placeholder = f"{{{key}}}"
            if placeholder in para.text:
                # 保留原有格式的替换
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(val))


def _replace_in_tables(doc, replace_map: dict):
    """替换文档中所有表格里的占位符"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in replace_map.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(val))


def _replace_all(doc, replace_map: dict):
    """替换文档中所有占位符"""
    _replace_in_paragraphs(doc, replace_map)
    _replace_in_tables(doc, replace_map)


def load_template(template_path: str) -> dict:
    """加载 Word 模板并验证存在性"""
    try:
        p = Path(template_path)
        if not p.exists():
            return {"success": False, "error": f"模板文件不存在: {template_path}"}
        if p.suffix.lower() not in (".docx",):
            return {"success": False, "error": f"不支持的模板格式: {p.suffix}，需要 .docx"}

        from docx import Document
        doc = Document(template_path)

        # 收集所有占位符
        placeholders = set()
        for para in doc.paragraphs:
            import re
            found = re.findall(r"\{(\w+)\}", para.text)
            placeholders.update(found)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found = re.findall(r"\{(\w+)\}", para.text)
                        placeholders.update(found)

        return {
            "success": True,
            "message": f"模板已加载，发现 {len(placeholders)} 个占位符",
            "placeholders": sorted(placeholders),
        }
    except ImportError:
        return {"success": False, "error": "python-docx 未安装，请执行: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def fill_fields(template_path: str, fields: dict, output_filename: str = "") -> dict:
    """用字段字典填充模板中所有占位符"""
    try:
        from docx import Document

        p = Path(template_path)
        if not p.exists():
            return {"success": False, "error": f"模板文件不存在: {template_path}"}

        doc = Document(template_path)
        _replace_all(doc, fields)

        out_name = output_filename or f"filled_{timestamp()}.docx"
        out_path = str(output_path(out_name))
        doc.save(out_path)

        return {
            "success": True,
            "message": f"文档已填充并保存",
            "files": [out_path],
        }
    except ImportError:
        return {"success": False, "error": "python-docx 未安装，请执行: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def save_document(doc, output_filename: str = "") -> dict:
    """保存 Document 对象到输出目录"""
    try:
        out_name = output_filename or f"document_{timestamp()}.docx"
        out_path = str(output_path(out_name))
        doc.save(out_path)
        return {"success": True, "message": "文档已保存", "files": [out_path]}
    except Exception as e:
        return {"success": False, "error": str(e)}


def batch_generate(
    template: str = "",
    data_source: str = "",
    output_pattern: str = "doc_{row}.docx",
    output_dir: str = "",
) -> dict:
    """批量生成文档：为数据源中每条记录生成一份文档"""
    try:
        from docx import Document

        if not template or not Path(template).exists():
            return {"success": False, "error": f"模板文件不存在: {template}"}

        # 加载数据源
        try:
            records = _load_data_source(data_source)
        except Exception as e:
            return {"success": False, "error": f"数据源加载失败: {str(e)}"}

        if not records:
            return {"success": False, "error": "数据源为空"}

        output_base = Path(output_dir) if output_dir else output_path("")
        ensure_dir(output_base)

        generated = []
        for idx, record in enumerate(records):
            doc = Document(template)
            _replace_all(doc, record)

            # 文件名变量替换
            filename = output_pattern
            for key, val in record.items():
                filename = filename.replace(f"{{{key}}}", str(val))
            filename = filename.replace("{row}", str(idx + 1)).replace("{index}", str(idx + 1))
            if not filename.endswith(".docx"):
                filename += ".docx"

            out_path = str(output_base / filename)
            doc.save(out_path)
            generated.append(out_path)

        return {
            "success": True,
            "message": f"批量生成完成: {len(generated)} 份文档",
            "files": generated,
        }
    except ImportError:
        return {"success": False, "error": "python-docx 未安装，请执行: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def merge_documents(
    template: str = "",
    data_source: str = "",
    output_pattern: str = "merged_document.docx",
    fields: dict = None,
) -> dict:
    """合并多个文档或字段到单份文档（复用 fill_template 入口）"""
    try:
        from docx import Document

        # 如果有数据源，按每条记录生成后合并为一份文档
        if data_source:
            try:
                records = _load_data_source(data_source)
            except Exception as e:
                return {"success": False, "error": f"数据源加载失败: {str(e)}"}

            if not records:
                return {"success": False, "error": "数据源为空"}

            master = Document()
            for idx, record in enumerate(records):
                if template and Path(template).exists():
                    doc = Document(template)
                    _replace_all(doc, record)
                    # 复制段落
                    for para in doc.paragraphs:
                        master.add_paragraph(para.text)
                    # 添加分页
                    if idx < len(records) - 1:
                        master.add_page_break()
                else:
                    # 无模板：直接写入记录
                    master.add_paragraph(json.dumps(record, ensure_ascii=False, indent=2))

            out_name = output_pattern if output_pattern else f"merged_{timestamp()}.docx"
            out_path = str(output_path(out_name))
            master.save(out_path)
            return {
                "success": True,
                "message": f"已合并 {len(records)} 条记录",
                "files": [out_path],
            }
        else:
            # 按字段填充
            return fill_template(template_path=template, fields=fields or {}, output_filename=output_pattern)
    except ImportError:
        return {"success": False, "error": "python-docx 未安装，请执行: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def fill_template(
    template: str = "",
    data_source: str = "",
    output_pattern: str = "filled_{index}.docx",
    fields: dict = None,
) -> dict:
    """填充模板（handler 入口）— 支持字段字典或 CSV/JSON 数据源"""
    fields = fields or {}

    # 如果有数据源，批量处理
    if data_source:
        return batch_generate(
            template=template,
            data_source=data_source,
            output_pattern=output_pattern,
        )

    # 单份填充
    if not template:
        return {"success": False, "error": "请提供模板路径或数据源"}
    return fill_fields(template_path=template, fields=fields, output_filename=output_pattern)
